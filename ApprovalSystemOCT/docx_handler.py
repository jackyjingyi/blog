import logging
from datetime import datetime
import docx
from docx import Document
from docx.table import Table
import timeit


class TableHandler:
    template_name = ""

    def __call__(self, *args, **kwargs):
        self.table = kwargs['table']

    def replace_text(self, text_old, text_new, tc=None, contain=False, replace_all=False):

        if tc:
            if text_old in tc.text.strip():
                for paragraph in tc.paragraphs:
                    for run in paragraph.runs:
                        if run.text == text_old:
                            run.text = text_new
                            break
        else:
            for row in self.table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text == text_old:
                                run.text = text_new
                                break


class ProjectTableHandler(TableHandler):
    """
    Requirement Template
    """
    template_name = "requirement"
    template_map = {'^project_name$': {'pos': (0, 7, 0, 0), 'default': '', 'default_replace': '', 'selection': False},
                    '^project_type1$': {'pos': (1, 1, 0, 0), 'default': '前瞻战略研究□', 'default_replace': '前瞻战略研究⌧',
                                        'selection': True},
                    '^project_type2$': {'pos': (1, 5, 0, 0), 'default': '产品管理研究□', 'default_replace': '产品管理研究⌧',
                                        'selection': True},
                    '^project_type3$': {'pos': (3, 7, 0, 0), 'default': '产品创新研究□', 'default_replace': '产品创新研究⌧',
                                        'selection': True},
                    '^p_rsd1$': {'pos': (2, 1, 0, 0), 'default': '新型城镇化□', 'default_replace': '新型城镇化⌧',
                                 'selection': True},
                    '^p_rsd2$': {'pos': (2, 4, 0, 0), 'default': '文化+品牌□', 'default_replace': '文化+品牌⌧',
                                 'selection': True},
                    '^p_rsd3$': {'pos': (2, 6, 0, 0), 'default': '主题公园+新场景□', 'default_replace': '主题公园+新场景⌧',
                                 'selection': True},
                    '^p_rsd4$': {'pos': (3, 0, 0, 0), 'default': '新商业□', 'default_replace': '新商业⌧', 'selection': True},
                    '^p_rsd5$': {'pos': (3, 3, 0, 0), 'default': '产品管理□', 'default_replace': '产品管理⌧',
                                 'selection': True},
                    '^p_rsd6$': {'pos': (3, 6, 0, 0), 'default': '旅游大数据□', 'default_replace': '旅游大数据⌧',
                                 'selection': True},
                    '^p_rsd7$': {'pos': (4, 0, 0, 0), 'default': '康旅□', 'default_replace': '康旅⌧', 'selection': True},
                    '^p_rsd8$': {'pos': (4, 3, 0, 0), 'default': '创新投□', 'default_replace': '创新投⌧', 'selection': True},
                    '^p_rsd9$': {'pos': (4, 6, 0, 0), 'default': '其它□', 'default_replace': '其它⌧', 'selection': True},
                    '^project_department$': {'pos': (5, 6, 0, 0), 'default': '', 'default_replace': '',
                                             'selection': False},
                    '^project_sponsor$': {'pos': (6, 1, 0, 0), 'default': '', 'default_replace': '',
                                          'selection': False},
                    '^project_phone$': {'pos': (6, 6, 0, 0), 'default': '', 'default_replace': '', 'selection': False},
                    '^project_co_group$': {'pos': (7, 6, 0, 0), 'default': '', 'default_replace': '',
                                           'selection': False},
                    '^project_funding1$': {'pos': (8, 1, 0, 0), 'default': '', 'default_replace': '',
                                           'selection': False},
                    '^project_funding2$': {'pos': (8, 6, 0, 0), 'default': '', 'default_replace': '',
                                           'selection': False},
                    '^project_start_time$': {'pos': (9, 6, 0, 1), 'default': '', 'default_replace': '',
                                             'selection': False},
                    '^project_end_time$': {'pos': (9, 6, 0, 3), 'default': '', 'default_replace': '',
                                           'selection': False},
                    '^project_detail$': {'pos': (10, 6, 0, 0), 'default': '', 'default_replace': '',
                                         'selection': False},
                    '^project_purpose$': {'pos': (11, 6, 1, 0), 'default': '', 'default_replace': '',
                                          'selection': False},
                    '^project_preparation$': {'pos': (12, 6, 1, 0), 'default': '', 'default_replace': '',
                                              'selection': False},
                    '^project_difficult$': {'pos': (13, 6, 1, 0), 'default': '', 'default_replace': '',
                                            'selection': False},
                    }

    def find_cell(self, target):
        for rdx, row in enumerate(self.table.rows):
            for cdx, cell in enumerate(row.cells):
                if target in cell.text.strip():
                    return (cell, rdx, cdx)

    def get_target_map(self, target):
        return self.template_map.get(target).get('pos')

    def get_target_run(self, map):
        rdx, cdx, pdx, rudx = map
        return self.table.rows[rdx].cells[cdx].paragraphs[pdx].runs[rudx]

    def replace_run_text(self, target, text, run):
        new_text = run.text.replace(target, text)
        run.text = new_text
        return run

    def get_default(self, target):
        return (self.template_map.get(target).get('default'), self.template_map.get(target).get('default_replace'))

    def fill_select(self, target, checked=False):
        map = self.get_target_map(target)
        run = self.get_target_run(map)
        t, r = self.get_default(target)
        if checked:
            return self.replace_run_text(target, r, run)
        else:
            return self.replace_run_text(target, t, run)

    def fill(self, target, text):
        map = self.get_target_map(target)
        run = self.get_target_run(map)
        return self.replace_run_text(target, text, run)

    def dispatch(self, value):
        # value is a dict from requirement.objects.values().get(pk=n)
        value['project_start_time'] =  datetime.strftime(value.get('project_start_time'),"%Y年%m月%d日")
        value['project_end_time'] = datetime.strftime(value.get('project_end_time'), "%Y年%m月%d日")

        _map = {'1': '^project_type1$', '2': '^project_type2$', '3': '^project_type3$'}
        _map_rsd = {
            '2': '^p_rsd1$',
            '3': '^p_rsd2$',
            '4': '^p_rsd3$',
            '5': '^p_rsd4$',
            '6': '^p_rsd5$',
            '7': '^p_rsd6$',
            '8': '^p_rsd7$',
            '9': '^p_rsd8$',
            '10': '^p_rsd9$'
        }
        _map_block = {
            'project_name': '^project_name$',
            'project_department': '^project_department$',
            'project_department_sponsor': '^project_sponsor$',
            'project_department_phone': '^project_phone$',
            'project_co_group': '^project_co_group$',
            'project_research_funding': '^project_funding1$',
            'project_outsourcing_funding': '^project_funding2$',
            'project_start_time': '^project_start_time$',
            'project_end_time': '^project_end_time$',
            'project_detail': '^project_detail$',
            'project_purpose': '^project_purpose$',
            'project_preparation': '^project_preparation$',
            'project_difficult': '^project_difficult$',
        }
        for key, val in value.items():
            try:
                if key == 'id':
                    continue
                if key == "project_type":
                    for k, v in _map.items():
                        if k == val:
                            self.fill_select(v, checked=True)
                        else:
                            self.fill_select(v)
                elif key == "project_research_direction":

                    for k1, v1 in _map_rsd.items():

                        if k1 in val:
                            self.fill_select(v1, checked=True)
                        else:
                            self.fill_select(v1)
                else:
                    self.fill(_map_block.get(key), val)
            except Exception as e:
                logging.warning(e)
                print(key, val)
                continue


if __name__ == '__main__':
    # document = Document(r'C:\Users\Jack\PycharmProjects\django_blog\media\template_doc1.docx')

    # tbl = document.tables
    print(__file__)
